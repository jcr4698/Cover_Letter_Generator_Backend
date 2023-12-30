import os
import sys

# docx modules
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from datetime import date

# cover letter backend modules
from generation.cover_letter_utils import add_hyperlink, get_cover_letter_text, get_body_text, generate_pdf, get_slash

class CoverLetter:
    
    def __init__(self):
        # internal variables
        self._curr_path = os.path.abspath(os.getcwd())
        os.chdir("./pdfs/")
        self._pdf_path = os.path.abspath(os.getcwd())
        os.chdir("../docx_files")
        self._docx_path = os.path.abspath(os.getcwd())
        os.chdir("..")
        self._MONTHS = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
        
        # external variables
        self._full_name = ""
        self._phone_num = ""
        self._email = ""
        self._linkedin = ""
        self._company = ""
        self._cover_body_ptr = None
        self._cover_body_txt = ""
    
    def set_full_name(self, full_name:str):
        self._full_name = full_name

    def get_full_name(self):
        return self._full_name
    
    def set_phone_num(self, phone_num:str):
        self._phone_num = phone_num

    def get_phone_num(self):
        return self._phone_num
    
    def set_email(self, email:str):
        self._email = email
    
    def get_email(self):
        return self._email
    
    def set_linkedin(self, linkedin:str):
        self._linkedin = linkedin
    
    def get_linkedin(self):
        return self._linkedin
    
    def set_company(self, company:str):
        self._company = company
    
    def get_company(self):
        return self._company
    
    def set_cover_letter_ptr(self, cover_body_filename:str):
        try:
            cover_letter_body_ptr = open(cover_body_filename, "r")
            self._cover_body_ptr = cover_letter_body_ptr
            return True
        except FileNotFoundError:
            print("'{}' not found.".format(cover_body_filename))
            return False
        
    def set_cover_letter_body(self, cover_body_text:str):
        self._cover_body_txt = cover_body_text
        
    def get_cover_letter_body(self):
        return self._cover_body_txt
    
    def close_cover_letter_ptr(self):
        self._cover_body_ptr
    
    def generate_cover_letter(self, filename:str):
        # Create Document
        document = Document()

        # Add Header
        cover_letter_header_details = {
            "name": self._full_name,
            "phone_num": self._phone_num,
            "email": self._email,
            "linkedin": self._linkedin
        }
        self._cover_letter_header_section(document, cover_letter_header_details)

        # Add Body
        curr_date = date.today()
        curr_month = self._MONTHS[curr_date.month-1]
        curr_day = curr_date.day
        curr_year = curr_date.year
        cover_letter_body_details = {
            "curr_date": "{} {}, {}".format(curr_month, curr_day, curr_year),
            "company": self._company,
            "name": self._full_name
        }
        body_txt = get_body_text(self._cover_body_txt, cover_letter_body_details)
        self._cover_letter_body_section(document, body_txt)

        # Document Margins
        for section in document.sections:
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)

        # Name or rename the docx
        file_no_type = os.path.splitext(" " + filename)[0].strip()
        file_rename = None
        if(file_no_type):
            file_rename = os.path.basename(file_no_type)
        else:
            file_rename = "auto_named_file"
            print("The given file name is invalid, so file has been automatically named '{}'.".format(file_rename))
            
        # Save the docx with new name
        if(file_rename):
            
            # Assign name to docx and pdf
            docx_filename = "{}{}".format(file_rename, ".docx")
            pdf_filename = "{}{}".format(file_rename, ".pdf")
            
            # Generate the docx file
            document.save(docx_filename)
            
            DIR_SLASH = get_slash(sys.platform)
            
            # Convert docx to pdf (only if docx saves successfully)
            curr_docx_path = "{}{}{}".format(self._curr_path, DIR_SLASH, docx_filename)
            dest_docx_path = "{}{}{}".format(self._docx_path, DIR_SLASH, docx_filename)
            curr_pdf_path = "{}{}{}".format(self._curr_path, DIR_SLASH, pdf_filename)
            dest_pdf_path = "{}{}{}".format(self._pdf_path, DIR_SLASH, pdf_filename)
            generate_pdf(curr_docx_path, dest_docx_path, curr_pdf_path, dest_pdf_path)
    
    ''' Header Section '''
    def _cover_letter_header_section(self, document:Document, cover_letter_details:dict):
        # initialize header
        header_section = document.sections[0].header

        # add and format header paragraph
        header = header_section.add_paragraph()
        remove_parag = header_section.paragraphs[0]._element
        remove_parag.getparent().remove(remove_parag)
        remove_parag._p = remove_parag._element = None
        header.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # add and format name to header paragraph
        header_name = header.add_run()
        header_name.text = cover_letter_details["name"]
        document.styles.add_style("Cover_Letter_Style_Header_Name", WD_STYLE_TYPE.CHARACTER)
        header_name_style = document.styles["Cover_Letter_Style_Header_Name"]
        header_name_style.font.name = "Calibri"
        header_name_style.font.size = Pt(13)
        header_name_style.font.bold = True
        header_name.style = header_name_style

        # add and format details to header paragraph
        header_details = header.add_run()
        header_details.text = "\n{} | ".format(cover_letter_details["phone_num"]) # phone number
        add_hyperlink(header, cover_letter_details["email"], cover_letter_details["email"]) # gmail link
        header_details = header.add_run()
        header_details.text = " | " # space
        add_hyperlink(header, cover_letter_details["linkedin"], cover_letter_details["linkedin"]) # linkedin link

        # set style for header details
        header_details_style = document.styles["Normal"]
        header_details_style.font.name = "Calibri"
        header_details_style.font.size = Pt(10)
        header_details_style.font.bold = False

    ''' Body Content Section '''
    def _cover_letter_body_section(self, document:Document, cover_letter_text:str):

        # add and format letter section
        letter = document.add_paragraph()

        # add and format name to introduction paragraph
        letter_content = letter.add_run()
        letter_content.text = cover_letter_text # insert cover letter text
        document.styles.add_style("Cover_Letter_Style_Content", WD_STYLE_TYPE.CHARACTER)
        letter_content_style = document.styles["Cover_Letter_Style_Content"]
        letter_content_style.font.name = "Calibri"
        letter_content_style.font.size = Pt(11)
        letter_content_style.font.bold = False
        letter_content.style = letter_content_style

    # def get_email(self):
    #     return self._email
    
# cv = CoverLetter()
# cv.set_full_name("Jan Carlos Rubio Sánchez")
# cv.set_phone_num("830-421-0344")
# cv.set_email("jcaj750@gmail.com")
# cv.set_linkedin("https://www.linkedin.com/in/jan-carlos-rubio-sanchez/")
# cv.set_company("TACC")
# cv.set_cover_letter_ptr("/mnt/c/Users/jan/Desktop/Files/Personal/Fun_Projects/Cover_Letter_Generator/cover_letters/cover_letter_body.txt")
# cv.generate_cover_letter()
# cv.close_cover_letter_ptr()