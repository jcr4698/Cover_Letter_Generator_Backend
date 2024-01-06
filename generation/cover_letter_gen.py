from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Inches
from cover_letter_utils import add_hyperlink, get_cover_letter_text, generate_pdf
from datetime import date

def main():
    # Create Document
    document = Document()

    # Add Header
    # cover_letter_header_details = {
    #     "name": "Jan Carlos Rubio Sánchez",
    #     "phone_num": "830-421-0344",
    #     "email": "jcaj750@gmail.com",
    #     "linkedin": "https://www.linkedin.com/in/jan-carlos-rubio-sanchez/"
    # }
    cover_letter_header_details = {
        "name": "Jan Carlos Rubio Sánchez",
        "phone_num": "830-421-0344",
        "email": "jcaj750@gmail.com",
        "linkedin": "https://www.linkedin.com/in/jan-carlos-rubio-sanchez/"
    }
    cover_letter_header_section(document, cover_letter_header_details)

    # Add Body
    months = ["January", "February", "March", "April", "May", "June", "July", "August", "September", "October", "November", "December"]
    curr_date = date.today()
    curr_month = months[curr_date.month-1]
    curr_day = curr_date.day
    curr_year = curr_date.year
    cover_letter_body_details = {
        "curr_date": "{} {}, {}".format(curr_month, curr_day, curr_year),
        "company": "Tag Strategies",
        "name": "Jan Carlos Rubio Sánchez"
    }
    cover_letter_body = get_cover_letter_text(cover_letter_body_details)
    cover_letter_body_section(document, cover_letter_body)

    # Document Margins
    for section in document.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # save the docx
    document.save("cover_letter_generation.docx")

    # convert docx to pdf
    generate_pdf("cover_letter_generation.docx")

''' Header Section '''
def cover_letter_header_section(document:Document, cover_letter_details:dict):
    # create header
    header_section = document.sections[0].header

    # add and format header paragraph
    header = header_section.add_paragraph()
    remove_parag = header_section.paragraphs[0]._element
    remove_parag.getparent().remove(remove_parag)
    remove_parag._p = remove_parag._element = None
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # add and format name to header paragraph
    header_name = header.add_run()
    header_name.text = cover_letter_details["name"]
    document.styles.add_style("Cover_Letter_Style_Header_Name", WD_STYLE_TYPE.CHARACTER)
    header_name_style = document.styles["Cover_Letter_Style_Header_Name"]
    header_name_style.font.name = "Calibri"
    header_name_style.font.size = Pt(13)
    header_name_style.font.bold = True
    header_name.style = header_name_style

    # add and format details to header paragraph
    # document.styles.add_style("Cover_Letter_Style_Header_Details", WD_STYLE_TYPE.CHARACTER)

    # add details to header details
    header_details = header.add_run()
    header_details.text = "\n{} | ".format(cover_letter_details["phone_num"]) # phone number
    add_hyperlink(header, cover_letter_details["email"], cover_letter_details["email"]) # gmail link
    header_details = header.add_run()
    header_details.text = " | " # space
    # linkedin_full_url = cover_letter_details["linkedin"]
    # linkedin_disp_url = linkedin_full_url[12: len(linkedin_full_url)-1]
    # add_hyperlink(header, linkedin_disp_url, linkedin_full_url) # linkedin link
    add_hyperlink(header, cover_letter_details["linkedin"], cover_letter_details["linkedin"]) # linkedin link

    # set style for header details
    # header_details_style = document.styles["Cover_Letter_Style_Header_Details"]
    header_details_style = document.styles["Normal"]
    header_details_style.font.name = "Calibri"
    header_details_style.font.size = Pt(10)
    header_details_style.font.bold = False

''' Body Content Section '''
def cover_letter_body_section(document:Document, cover_letter_text:str):
    # add and format letter section
    letter = document.add_paragraph()

    # add and format name to introduction paragraph
    letter_content = letter.add_run()
    letter_content.text = cover_letter_text # insert cover letter text
    document.styles.add_style("Cover_Letter_Style_Content", WD_STYLE_TYPE.CHARACTER)
    letter_content_style = document.styles["Cover_Letter_Style_Content"]
    letter_content_style.font.name = "Calibri"
    letter_content_style.font.size = Pt(11)
    letter_content_style.font.bold = False
    letter_content.style = letter_content_style

if(__name__ == "__main__"):
    main()